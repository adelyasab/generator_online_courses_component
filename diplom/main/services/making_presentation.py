from docx import Document


def extract_images_from_docx(docx_file):
    doc = Document(docx_file)
    extracted_images = []

    for rel in doc.part.rels.values():
        print(rel.reltype)
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            extracted_images.append(image_data)

    return extracted_images


# Пример использования:
docx_file = "text_for_presentation.docx"
images = extract_images_from_docx(docx_file)

for i, image in enumerate(images, start=1):
    with open(f"image_{i}.jpg", "wb") as f:
        f.write(image)


def extract_bold_text(docx_file):
    doc = Document(docx_file)
    bold_text = []

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.bold:
                bold_text.append(run.text)

    return bold_text


# Пример использования:
docx_file = 'text_for_presentation.docx'
bold_text = extract_bold_text(docx_file)
print(bold_text)
